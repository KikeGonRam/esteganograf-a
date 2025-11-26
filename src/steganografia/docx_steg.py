from docx import Document
import base64
import os

class DocxSteganography:
    @staticmethod
    def encode(carrier_path, out_path, payload, password=None):
        # Oculta el payload como comentario en el DOCX
        doc = Document(carrier_path)
        # Insertar como comentario oculto
        doc.add_paragraph(f"<!--hidden_payload:{base64.b64encode(payload.encode('utf-8')).decode('utf-8')}-->")
        doc.save(out_path)

    @staticmethod
    def decode(carrier_path, password=None):
        doc = Document(carrier_path)
        for para in doc.paragraphs:
            text = para.text
            if '<!--hidden_payload:' in text:
                start = text.find('<!--hidden_payload:') + len('<!--hidden_payload:')
                end = text.find('-->', start)
                if end != -1:
                    b64 = text[start:end]
                    return base64.b64decode(b64).decode('utf-8')
        raise Exception('No se encontró información oculta en el DOCX')
